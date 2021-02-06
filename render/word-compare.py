import sys
import docx
from win32com import client
import os
import requests
import json

def getpath():
    file=sys.argv
    return file

def readdocx(path):
    #读取文档内容，返回文档数据
    file=docx.Document(path)
    tables=[]
    #检验这个文件的表是否我们检查的表
    for item in file.tables:
        count=0
        for i in item.rows:
            rows={}
            if count==0:
                for ex in i.cells:
                    rows[ex.text]=ex.text
            else:
                cell_count=0
                for ex in i.cells:
                    rows[item.rows[0].cells[cell_count].text]=ex.text
                    cell_count+=1
            tables.append(rows)
            count+=1
    
    return tables

def compare_docx(main_data):#校验文档
    filelist=getpath()
    main_key=["标准名称","标准号","发布日期","实施日期","状态"]
    cmp_result=[]
    for item in range(1,len(filelist)):
        docx_data=readdocx(filelist[item])#比较这个文件是否有和主库相关的内容，没有则不解析
        if len(docx_data)<1:
            print("100")
            continue 
        flag=False
        for a in main_key:
            for i in docx_data[0]:
                if a==docx_data[0][i]:
                    flag=True
                    break
            if flag==True:
                break
        if flag==False:
            print("100")
            continue  
        for i in docx_data:
            rows={}
            for ex in i:
                cmp_key=""
                for a in main_key:#看这个docx_data里关键字主库里有没有，如果有则进一步比较
                    if ex==a:
                        cmp_key=ex
                        break
                if cmp_key:#比较在这个关键字下是否找到和主库相同的信息,如果找到则在
                    for a in main_data:
                        if a[cmp_key]==i[ex]:
                            rows[cmp_key]="正确"
                            break
                    if cmp_key not in rows:
                        rows[cmp_key]="错误"
                else:
                    rows[ex]="这个列项不是主库内的内容"
            cmp_result.append(rows)
        
        writeresult(docx_data,cmp_result,filelist[item])

def writeresult(docx_data,cmp_result,path):#将比较结果写成docx文档
    document=docx.Document()
    table=document.add_table(1,len(docx_data[1]))
    table.autofit=True
    for i in range(0,len(docx_data)):

        if i==0:
            count=0
            for a in docx_data[1]:
                table.rows[0].cells[count].text=a
                count+=1
            
        else:
            table.add_row()
            table.add_row()
            item=i*2
            count=0
            for a in docx_data[i]:
                table.rows[item-1].cells[count].text=docx_data[i][a]
                count+=1          
            count=0
            for a in cmp_result[i]:
                table.rows[item].cells[count].text=cmp_result[i][a]
                count+=1
    #再将数据打包成json文件，供软件呈现
    json_data=[]
    for i in range(1,len(docx_data)):
        rows={}
        for ex in docx_data[i]:
            rows[ex]=docx_data[i][ex]
        json_data.append(rows)
        rows={}
        for ex in cmp_result[i]:
            rows[ex]=cmp_result[i][ex]
        json_data.append(rows)
    json_data=json.dumps(json_data)
    
    path=path.split("\\")
    json_data_url=os.getcwd()+"\\..\\data\\"
    str_name=path[len(path)-1]
    str_name=str_name.replace(".docx",".json")
    str_name="result_{}".format(str_name)
    json_data_url+=str_name
    with open(json_data_url,"w",encoding="utf-8") as f:
        f.write(json_data)
    path[len(path)-1]="result_{}".format(path[len(path)-1])
    path="\\".join(path)
    document.save(path)


def readdata():#读取主库，与主库相比
    url=os.getcwd()+"\\..\\data\\a.json"
    res=open(url,"r",encoding="utf-8")
    res=json.load(res)
    return res


if __name__ == "__main__":
    compare_docx(readdata())
    print("200")